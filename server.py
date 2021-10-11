from http.server import HTTPServer, BaseHTTPRequestHandler
import json
import os
from glob import glob
from pathlib import Path
from urllib.parse import unquote
from scipy.signal import savgol_filter, argrelextrema, find_peaks, resample_poly
from scipy.optimize import differential_evolution, curve_fit
import numpy
import pandas
#Also requires openpyxl
import zipfile
import tempfile
import shutil
import xlsxwriter
import webbrowser
from multiprocessing import Pool, freeze_support
import warnings
import math
#from adjustText import adjust_text

#NEW
from matplotlib import pyplot
import matplotlib
from docx import Document
import argparse
import configparser

np =numpy

#/Users/pau/HTS/
serv_path="webapp/"
data_path="data/"
cache_path="cache/"
persistent_path="persistent/"
plateinfo_path="plateinfo/"
version="7"
#decimation=False 
decimation=True
decim_maxpoints=150

formats=["LightCycler","BioRad","QuantStudio 2","QuantStudio 1","Generic DSF"]



def string_escape(s, encoding='utf-8'):
    return (s.encode('latin1')         # To bytes, required by 'unicode-escape'
             .decode('unicode-escape') # Perform the actual octal-escaping decode
             .encode('latin1')         # 1:1 mapping back to bytes
             .decode(encoding))  

def remove_cache():
    files = glob(cache_path+'/*')
    for f in files:
        os.remove(f)


def kd_function(x,kd,dH0,Tm0):
    R=8.314
    n=1
    #Tm0=348.98

    f=Tm0/(1-n*R*Tm0/dH0*np.log(1+x/kd))
    
    #f=dH0*Tm0/ (dH0-R*Tm0*np.log((kd+x)/kd))
    
    return f

def sumOfSquaredError(parameterTuple,func,xdata,ydata):
    warnings.filterwarnings("ignore") # do not print warnings by genetic algorithm
    #print(parameterTuple,xdata,ydata)
    val = func(xdata, *parameterTuple)
    return np.sum((ydata - val) ** 2.0)



#https://stackoverflow.com/questions/56164659/scipy-curve-fit-do-not-converge-even-if-i-iteratively-change-initial-guess
def generate_Initial_Parameters(func,xdata,ydata):
    
    #finds best starting parameters to fit xdata and ydata to func.
    minY = np.min(ydata)
    maxY = np.max(ydata)
    minX = np.min(xdata)
    maxX = np.max(xdata)

    parameterBounds = []
    parameterBounds.append([minX, maxX]) # search bounds for kd
    parameterBounds.append([-1e8, 1e8]) # search bounds for dh0
    parameterBounds.append([minY, maxY]) # search bounds for tm

    # "seed" the numpy random number generator for repeatable results
    result = differential_evolution(sumOfSquaredError, parameterBounds, args=([func,xdata,ydata]),seed=3)
    return result.x

def getKd(concentrations,tms):
    #We need the data at x and y.
    #file="AT021-D02.txt"
    #df=read_csv(file,sep="\t")
    #data=df.values



    #x,y=data[:,0], data[:,1]
    x,y=concentrations,tms
    minY = np.min(y)
    maxY = np.max(y)
    minX = np.min(x)
    maxX = np.max(x)

    # by default, differential_evolution completes by calling curve_fit() using parameter bounds
    geneticParameters = generate_Initial_Parameters(kd_function,x,y)
    #print(geneticParameters)
    #geneticParameters.append(maxfev=100000)
    # now call curve_fit without passing bounds from the genetic algorithm,
    # just in case the best fit parameters are aoutside those bounds

    b11=0.9*geneticParameters[0]
    b12=1.1*geneticParameters[0]

    if (b12>maxX):
        b12=maxX


    b21=0.9*geneticParameters[1]
    b22=1.1*geneticParameters[1]
    if geneticParameters[1]<0:
        b21=1.1*geneticParameters[1]
        b22=0.9*geneticParameters[1]

    b31=0.9*geneticParameters[2]
    b32=1.1*geneticParameters[2]

    try:
        #kd,th0,tm0
        #fittedParameters, pcov = curve_fit(kd_function, x, y, p0=geneticParameters, method='trf',maxfev=100000, bounds=([minX, -1e12,0.8*minY],[maxX,1e12,1.5*maxY]))
        #fittedParameters, pcov = curve_fit(kd_function, x, y, p0=geneticParameters, method='dogbox',maxfev=100000, bounds=([minX, -1e8,0.8*minY],[maxX,1e8,1.5*maxY]))
        #fittedParameters, pcov = curve_fit(kd_function, x, y, p0=geneticParameters, method='lm',maxfev=100000)#, bounds=([minX, -1e8,0.8*minY],[maxX,1e8,1.5*maxY]))
        fittedParameters, pcov = curve_fit(kd_function, x, y, p0=geneticParameters, method='trf',maxfev=100000, bounds=([b11,b21,b31],[b12,b22,b32]))

    except:
        #print ("Kd calculation not possible with these values")
        return(0,0,"ERROR","ERROR","ERROR")
    #print('Fitted parameters:', fittedParameters)
    #print('pcov',pcov)

    modelPredictions = kd_function(x, *fittedParameters) 

    absError = modelPredictions - y

    SE = np.square(absError) # squared errors
    MSE = np.mean(SE) # mean squared errors
    RMSE = np.sqrt(MSE) # Root Mean Squared Error, RMSE
    Rsquared = 1.0 - (np.var(absError) / np.var(y))
    
    tkd="{:.2e}".format(fittedParameters[0])

    if (fittedParameters[0]>1):
        tkd="{:.2f}".format(fittedParameters[0])
    tdH0="{:.2f}".format(fittedParameters[1])
    ttm0="{:.2f}".format(fittedParameters[2])
    #print('Kd: '+tkd+" M ; dH0: "+ tdH0+" J ; Tm0: "+ttm0+ " K")
    #print('RMSE:', RMSE)
    #print('R-squared:', Rsquared)



    return(fittedParameters,Rsquared,tkd,tdH0,ttm0)

def get_kd_data_for_molecule(template,plate,tm_ave,molecule):
    #print("getting Kd data for molecule",molecule,template,plate,tm_ave)
    output_data={}
    #output data contains:{wells:{'A1':{'tm':65,'enabled':1,'concentration':250},'A2':...},'datapoints':[{"x":12,"y:7"},{"x":12,"y:7"},...],'kd':0.55,'R2':0.999,'dH0':-544445,'tm0':312,'fittingcurve':[{"x":12,"y:7"},{"x":12,"y:7"},...]}
    kd_templatefile=persistent_path+"kdtemplates.templ"
    templ={}
    if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            for item in a:
                #print(item)
                #print(item["filename"],template)
                if item["filename"]==template:
                    templ=item
    moleculewells=[]
    for item in templ["cells"]:
        #print(item)
        
        for key in item:
            #print(key,item[key])
            splkey=key.split("-")
            w=chr(int(splkey[1])+64)+str(splkey[0])
            molecul=item[key]
            if (molecul==molecule):
                moleculewells.append(w)
            #print(w)
    

    concentrations={}
    for item in templ["concentrations"]:
        for key in item:
            cl=key[1:]
            a=item[key]
            if a=="":
                a="0"
            conc=float(a)
            concentrations[cl]=conc
    #print(concentrations)

    wells={}
    tmdata=readtxttm(data_path+plate)
    cats_item=json.loads(serve_categories(plate))

    tmconc0=0
    tmconc50=0
    tmconcmax=0
    for well in moleculewells:
        good_data=1
        if well in cats_item:
            dat=cats_item[well]
            if ("Warning") in dat:
                good_data=0
        tm=float(tmdata[well]["tm"])+273.15
        allpeaks=[]
        if 'allpeaks' in tmdata[well]:
            allpeaks=tmdata[well]['allpeaks']
        if len(allpeaks)>1:
            #closest value to tm_ave
            tm=allpeaks[numpy.abs(numpy.array(allpeaks)-tm_ave).argmin()]+273.15

        concentration=concentrations[well[1:]]
        wells[well]={'tm':tm,'enabled':good_data,'concentration':concentration}
    
    output_data["wells"]=wells
    datapoints=[]
    concs_for_kd=[]
    tms_for_kd=[]

    for item in wells:
        datapoints.append({'x':wells[item]['concentration'],'y':wells[item]['tm'],'label':item,'enabled':wells[item]['enabled'],'well':item,'molecule':molecule})
        if (wells[item]['enabled']==1):
            concs_for_kd.append(wells[item]['concentration'])
            tms_for_kd.append(wells[item]['tm'])

    output_data["datapoints"]=datapoints

    fittedParameters,Rsquared,tkd,tdH0,ttm0=getKd(concs_for_kd,tms_for_kd)
    output_data['fkd']=tkd
    output_data['kd']=fittedParameters[0]
    output_data['dH0']=tdH0
    output_data['tm0']=ttm0
    output_data['R2']=Rsquared
    output_data['max_conc']=np.max(concs_for_kd)
    output_data['unfinished_curve']=0
    if (abs(fittedParameters[0]-np.max(concs_for_kd))<0.001):
        output_data['unfinished_curve']=1
    
    #Now generate fitting curve:
    output_data['fittingcurve']=[]
    xf=[]
    yf=[]
    if (Rsquared>0):
        xf=np.linspace(np.min(concs_for_kd),np.max(concs_for_kd),100)
        yf=kd_function(xf,fittedParameters[0],fittedParameters[1],fittedParameters[2])
    

    for i in range(0,len(xf)):
        output_data['fittingcurve'].append({'x':xf[i],'y':yf[i],'label':''})
    #print(xf)
    #print(yf)

    
    #print("\n")
    #print(output_data)
    return output_data

def get_kd_data_m(post_data):
    output_data={}
    p=json.loads(post_data)
    filename=p["filename"]
    template=p["template"]
    refs=p["references"]
    manual_ref=p["manual_ref"]
    molecule=p["molecule"]
    tm_ave,tm_std=calculate_ref_tm(filename,refs,manual_ref)
    output_data[molecule]=get_kd_data_for_molecule(template,filename,tm_ave,molecule)
    return output_data

def get_kd_data(post_data):
    
    output_data={}
    p=json.loads(post_data)
    filename=p["filename"]
    template=p["template"]
    refs=p["references"]
    manual_ref=p["manual_ref"]
    

    #Load template:
    kd_templatefile=persistent_path+"kdtemplates.templ"
    templ={}
    if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            for item in a:
                #print(item)
                #print(item["filename"],template)
                if item["filename"]==template:
                    templ=item
    moleculewells={"A":[],"B":[],"C":[],"D":[],"E":[],"F":[],"G":[],"H":[],"I":[],"J":[],"K":[],"L":[],"M":[],"N":[],"O":[],"P":[],"Q":[],"R":[],"S":[],"T":[],"U":[],"V":[],"W":[],"X":[],"Y":[],"Z":[]}

    for item in templ["cells"]:
        #print(item)
        
        for key in item:
            #print(key,item[key])
            splkey=key.split("-")
            w=chr(int(splkey[1])+64)+str(splkey[0])
            molecule=item[key]
            if (molecule!=""):
                moleculewells[molecule].append(w)
    unique_mols=set()
    for item in moleculewells:
        if len(moleculewells[item])>0:
            unique_mols.add(item)
    unique_mols=sorted(list(unique_mols))

    #For each molecule in template, get_kd_data_for_molecule(template,plate,references,manual_ref,molecule)
    tm_ave,tm_std=calculate_ref_tm(filename,refs,manual_ref)
    
    for molec in unique_mols:
        output_data[molec]=get_kd_data_for_molecule(template,filename,tm_ave,molec)

    #print(output_data)
    return(output_data)

def calculate_ref_tm(plate,references,manual_ref):
    filename=plate
    refs=references
    temps_refs=[]
    tmdata=readtxttm(data_path+filename)
    cats_item=json.loads(serve_categories(filename))
        #CAlculate ave_tM
    if (manual_ref!=0):
        tm_ave=float(manual_ref)
        tm_std=0
    else:
        for rc in refs:
            #print(rc)
            splrc=rc.split()
            #print(splrc)
            if (splrc[0]=="Col"):
                colnum=int(splrc[1])
                for well in tmdata:
                    #print(well,well[:1],well[1:])
                    if (well[:1]!="d"):
                        #print(well,well[:1],well[1:],str(colnum))
                        if (well[1:]==str(colnum)):
                            #print(well)
                            good_data=True
                            if well in cats_item:
                                dat=cats_item[well]
                                if ("Warning") in dat:
                                    good_data=False
                                    #print("Bad reference at "+well+" "+item)
                            if (good_data):    
                                temps_refs.append(float(tmdata[well]["tm"]))
            elif (splrc[0]=="Row"):
                #print("ROW")
                rowletter=splrc[1]
                for well in tmdata:
                    #print(well,well[:1],well[1:])
                    if (well[:1]!="d"):
                        #print(well,well[:1],well[1:],str(rowletter))
                        if (well[:1]==str(rowletter)):
                            good_data=True
                            if well in cats_item:
                                dat=cats_item[well]
                                if ("Warning") in dat:
                                    good_data=False
                                    #print("Bad reference at "+well+" "+item)
                            if (good_data):    
                                temps_refs.append(float(tmdata[well]["tm"]))
            else:

                #Single cell.
                #print ("CELL")
                well=rc
                good_data=True
                if well in cats_item:
                    dat=cats_item[well]
                    if ("Warning") in dat:
                        good_data=False
                        #print("Bad reference at "+well+" "+item)
                if (good_data):    
                    temps_refs.append(float(tmdata[well]["tm"]))

        #print(temps_refs)
        tm_ave=numpy.average(temps_refs)
        tm_std=numpy.std(temps_refs)
    return(tm_ave,tm_std)


def create_kd_report(post_data):

    matplotlib.use('agg')
    

    tm_ave=0

    report_path=tempfile.mkdtemp()
    report_path_fil=report_path+os.path.sep+'report.docx'
    
    document= Document()
    
    #print(report_path_fil)
    p=json.loads(post_data)
    filename=p["filename"]
    template=p["template"]
    refs=p["references"]
    manual_ref=p["manual_ref"]
    temps_refs=[]
    #Calculate ref_temp to choose best tm maximum for multimaxima wells

            
    
    document.add_heading('KD Report for plate '+filename,0)
    document.add_page_break()
    #print(filename,template)

    #Open template:
    kd_templatefile=persistent_path+"kdtemplates.templ"
    templ={}
    if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            for item in a:
                #print(item)
                #print(item["filename"],template)
                if item["filename"]==template:
                    templ=item

    #print(templ)
    moleculewells={"A":[],"B":[],"C":[],"D":[],"E":[],"F":[],"G":[],"H":[],"I":[],"J":[],"K":[],"L":[],"M":[],"N":[],"O":[],"P":[],"Q":[],"R":[],"S":[],"T":[],"U":[],"V":[],"W":[],"X":[],"Y":[],"Z":[]}
    concentrations={"1":0,"2":0,"3":0,"4":0,"5":0,"6":0,"7":0,"8":0,"9":0,"10":0,"11":0,"12":0,"13":0,"14":0,"15":0,"16":0,"17":0,"18":0,"19":0,"20":0,"21":0,"22":0,"23":0,"24":0}

    for item in templ["cells"]:
        #print(item)
        
        for key in item:
            #print(key,item[key])
            splkey=key.split("-")
            w=chr(int(splkey[1])+64)+str(splkey[0])
            molecule=item[key]
            if (molecule!=""):
                moleculewells[molecule].append(w)
            #print(w)

    for item in templ["concentrations"]:
        for key in item:
            cl=key[1:]
            a=item[key]
            if a=="":
                a="0"
            conc=float(a)
            concentrations[cl]=conc




    #Now, get Tms from filename.
    tmdata=readtxttm(data_path+filename)
    #Calculate Tm and stdev for references.
    cats_item=json.loads(serve_categories(filename))

    #Get tm_Ave and tm_std
    tm_ave,tm_std=calculate_ref_tm(filename,refs,manual_ref)


    #print(moleculewells)
    #print(concentrations)
    for item in moleculewells:
        #For each molecule
        concs_for_kd=[]
        tms_for_kd=[]
        wells_for_kd=[]
        if len(moleculewells[item])>0:
            #print ("Calculating molecule "+ item)
            document.add_heading('Molecule '+item, level=1)
            for well in moleculewells[item]:
                good_data=True
                concentration=concentrations[well[1:]]
                if well in cats_item:
                        dat=cats_item[well]
                        if ("Warning") in dat:
                            good_data=False
                tm=float(tmdata[well]["tm"])+273.15


                allpeaks=[]
                if 'allpeaks' in tmdata[well]:
                    allpeaks=tmdata[well]['allpeaks']
                if len(allpeaks)>1:
                    #closest value to tm_ave
                    tm=allpeaks[numpy.abs(numpy.array(allpeaks)-tm_ave).argmin()]+273.15

                #print(well,tm,concentration,good_data)
                if (good_data):
                    concs_for_kd.append(concentration)
                    tms_for_kd.append(tm)
                    wells_for_kd.append(well)
            #print ("Molecule ",item, " ; wells ",moleculewells[item]) 
            #print("CONCS:",concs_for_kd)
            #print("TMS:",tms_for_kd)
            fittedParameters,Rsquared,tkd,tdH0,ttm0=getKd(concs_for_kd,tms_for_kd)
            #print("Molecule:",item," KD:",tkd, " dH0: ",tdH0," tm0: ",ttm0, 'R2: ',Rsquared)
            t=""
            for item2 in moleculewells[item]:
                if (t!=""):
                    t=t+" ,"
                t=t+item2
            document.add_heading("Wells: ",level=3)
            document.add_paragraph(t)
            if (abs(fittedParameters[0]-np.max(concs_for_kd))<0.001):
                tkd=">"+tkd
            document.add_heading("Fitting results:",level=3)
            document.add_paragraph("Kd: "+str(tkd) + "uM")
            document.add_paragraph("dH0: "+str(tdH0)+"J")
            document.add_paragraph("tm0: "+str(ttm0)+ 'K')
            document.add_paragraph("R2: "+str(Rsquared))
            #print(fittedParameters[0],np.max(concs_for_kd),abs(fittedParameters[0]-np.max(concs_for_kd)))
            if (abs(fittedParameters[0]-np.max(concs_for_kd))<0.001):
                document.add_heading("Unfinished curve. Kd is higher than maximum used concentration. Please use higher concentrations.",level=3)
            
                #Plot:
            x,y=concs_for_kd,tms_for_kd
            points=[]
            for i in range(0,len(x)):
                points.append((x[i],y[i],wells_for_kd[i]))
            pointss=sorted(points,key=lambda x:x[0])
            #print(pointss)

            #Set different color depending on wells_for_kd first letter.
            px,py,pl=[],[],[]
            for i in range(0,len(x)):
                px.append(pointss[i][0])
                py.append(pointss[i][1])
                pl.append(pointss[i][2])
            pyplot.ioff()
            pyplot.xscale('log')

            points_for_row={}
            for x,y,s in zip(px,py,pl):
                row=s[0]
                if row in points_for_row:
                    points_for_row[row].append((x,y,s))
                else:
                    points_for_row[row]=[(x,y,s)]

            #print (points_for_row)
            rows_set=set()
            for it in points_for_row:
                #print (it)
                rows_set.add(it)
            rows_list=sorted(list(rows_set))
            

            num_of_rows=len(rows_list)
            colors=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf','#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf','#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf','#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf']
            colors=['#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4','#1f77b4']
            ccount=0
            for it in rows_list:
                #print(points_for_row[it])
                ppx,ppy=[],[]
                for it2 in points_for_row[it]:
                    #print (it2)
                    ppx.append(it2[0])
                    ppy.append(it2[1])
                    pyplot.scatter(ppx,ppy,c=colors[ccount],marker=".")
                ccount=ccount+1


            #Write labels
            #texts=[]
            #for x,y,s in zip(px,py,pl):
            #    texts.append(pyplot.text(x,y,s))
            #adjust_text(texts, only_move={'points':'y', 'texts':'y'}, arrowprops=dict(arrowstyle="->", color='r', lw=0.5))

            #Point annotation
            #for i, txt in enumerate(pl):
            #    pyplot.annotate(txt, (px[i], py[i]))

            if (Rsquared>0):
                Kd,dH,T0=fittedParameters

                x_line = px
                y_line= kd_function(x_line,Kd,dH,T0)

                pyplot.plot(x_line,y_line)
            
            pyplot.savefig(report_path+os.path.sep+"molecule"+item+".png",format='png')
            document.add_picture(report_path+os.path.sep+"molecule"+item+".png")
            pyplot.clf()
            document.add_page_break()

            #pyplot.show()
    #Now, generate the tables at the end of the document:
    document.add_heading('Data for plate '+filename,0)
    for item in moleculewells:
        #For each molecule
        concs_for_kd=[]
        tms_for_kd=[]
        wells_for_kd=[]
        if len(moleculewells[item])>0:
            #print ("Calculating molecule "+ item)
            document.add_heading('Molecule '+item, level=1)
            for well in moleculewells[item]:
                good_data=True
                concentration=concentrations[well[1:]]
                if well in cats_item:
                        dat=cats_item[well]
                        if ("Warning") in dat:
                            good_data=False
                tm=float(tmdata[well]["tm"])+273.15


                allpeaks=[]
                if 'allpeaks' in tmdata[well]:
                    allpeaks=tmdata[well]['allpeaks']
                if len(allpeaks)>1:
                    #closest value to tm_ave
                    tm=allpeaks[numpy.abs(numpy.array(allpeaks)-tm_ave).argmin()]+273.15

                #print(well,tm,concentration,good_data)
                if (good_data):
                    concs_for_kd.append(concentration)
                    tms_for_kd.append(tm)
                    wells_for_kd.append(well)
            
        #Make table for molecule
        number_of_wells=len(wells_for_kd)
        if (number_of_wells)>0:
            numcols=8
            numrows=(int( (number_of_wells-1)/(numcols-1) )+1)*3
            #print("numcols",numcols,"numrows",numrows)
            table = document.add_table(rows=numrows, cols=numcols)
            table.style="TableGrid"

            for i in range(0,number_of_wells):
                c="{:.2e}".format(concs_for_kd[i])
                if (concs_for_kd[i]>1):
                    c="{:.2f}".format(concs_for_kd[i])
                row_num=(int(i/(numcols-1))*3)
                col_num=(i % (numcols-1))+1
                #print(row_num,col_num,wells_for_kd[i])
                #print("row_num",row_num,"col_num",col_num,"i",i)
                table.rows[0 + row_num].cells[col_num].text=wells_for_kd[i]
                table.rows[1 + row_num].cells[col_num].text=c
                table.rows[2 + row_num].cells[col_num].text="{:.3f}".format(tms_for_kd[i])
                table.rows[0 + row_num].cells[0].text="Well"
                table.rows[1 + row_num].cells[0].text="Conc."
                table.rows[2 + row_num].cells[0].text="Tm"

            document.add_page_break()

    document.save(report_path_fil)
    #f=open(report_path_fil,"w")
    #f.write(filename)
    #f.write(template)
    #f.close()
    return report_path_fil

def create_report(post_data):

    #Bad data  ("Warning") and "Validated" are exclusive
    report_path=tempfile.mkdtemp()
    report_path_fil=report_path+os.path.sep+'report.xlsx'
    #print(report_path_fil)
    report=xlsxwriter.Workbook(report_path_fil)
    worksheet1 = report.add_worksheet("Validated")
    worksheet2 = report.add_worksheet("Effect")
    worksheet3 = report.add_worksheet("All")
    worksheet4 = report.add_worksheet("Data")

    header_format= report.add_format()
    header_format.set_bold()

    neutral_format=report.add_format()
    neutral_format.set_bg_color("#FFFFFF")
    neutral_format.set_border(1)
    stab_format=report.add_format()
    stab_format.set_bg_color("#D5F5E3")
    stab_format.set_border(1)
    destab_format=report.add_format()
    destab_format.set_bg_color("#FDEBD0")
    destab_format.set_border(1)
    uncert_format=report.add_format()
    uncert_format.set_bg_color("#FDD0F5")
    uncert_format.set_border(1)

    report_data=[]
    #print(post_data)
    p=json.loads(post_data)
    #print(p)
    p_filenames=p["filenames"]
    threshold=p["threshold"]
    refs=p["references"]
    manual_ref=p["manual_ref"]
    spk_counter=1
    #print(p)



    #Read template database.
    templatesdf=[]
    #for path in Path(data_path).rglob(('*.txt','*.xlsx')):
    extensions=['*.txt']
    all_files = []
    for ext in extensions:
        all_files.extend(Path(plateinfo_path).rglob(ext))
    for filename in all_files:
            #print(filename)
            df = pandas.read_csv(filename, index_col=None, sep="\t",header=0)
            templatesdf.append(df)
    templatesdf=pandas.concat(templatesdf, axis=0, ignore_index=True)
    templatesdf.columns=["platewell","eos","smiles","plate","well"]
    #Columns used: 1,2,3,4
    #Columns ignored: 0
    #Skips header.


    for item in p_filenames:
        #print(item)
        
        
        template=serve_template(item)
        b=readtxttm(data_path+item)
        #Calculate Tm and stdev for references.
        cats_item=json.loads(serve_categories(item))
        #print(refs)
        temps_refs=[]
        if (manual_ref!=0):
            tm_ave=float(manual_ref)
            tm_std=0
        else:
            for rc in refs:
                #print(rc)
                splrc=rc.split()
                #print(splrc)
                if (splrc[0]=="Col"):
                    colnum=int(splrc[1])
                    for well in b:
                        #print(well,well[:1],well[1:])
                        if (well[:1]!="d"):
                            #print(well,well[:1],well[1:],str(colnum))
                            if (well[1:]==str(colnum)):
                                #print(well)
                                good_data=True
                                if well in cats_item:
                                    dat=cats_item[well]
                                    if ("Warning") in dat:
                                        good_data=False
                                        #print("Bad reference at "+well+" "+item)
                                if (good_data):    
                                    temps_refs.append(float(b[well]["tm"]))
                elif (splrc[0]=="Row"):
                    #print("ROW")
                    rowletter=splrc[1]
                    for well in b:
                        #print(well,well[:1],well[1:])
                        if (well[:1]!="d"):
                            #print(well,well[:1],well[1:],str(rowletter))
                            if (well[:1]==str(rowletter)):
                                good_data=True
                                if well in cats_item:
                                    dat=cats_item[well]
                                    if ("Warning") in dat:
                                        good_data=False
                                        #print("Bad reference at "+well+" "+item)
                                if (good_data):    
                                    temps_refs.append(float(b[well]["tm"]))
                else:

                    #Single cell.
                    #print ("CELL")
                    well=rc
                    good_data=True
                    if well in cats_item:
                        dat=cats_item[well]
                        if ("Warning") in dat:
                            good_data=False
                            #print("Bad reference at "+well+" "+item)
                    if (good_data):    
                        temps_refs.append(float(b[well]["tm"]))

            temps_refs=[i for i in temps_refs if i> 1]
            print(temps_refs)
            tm_ave=numpy.average(temps_refs)
            tm_std=numpy.std(temps_refs)
        #print(tm_ave,tm_std)
        
        
        for well in b:
                if well[0]=="d":
                    continue
                #print(well)
                if well=="format":
                    continue
                wd={"File":item,"Well":"","Validated":False,"Effect":"Neutral","Warning":False,"tm":0, "dtm":0,"ref":0,"ref_std":0,"spkline1":spk_counter,"spkline2":spk_counter+1,"num_datum":10,"extra_tm_peaks":"","template":template,"eos":"","smiles":"","absdtm":0}
                
                #Add leading zero to well
                w3=well
                if (len(w3)==2):
                    w3=w3[0]+"0"+w3[1]
                wd["Well"]=well
                #First, calculate values

                #tm may be the best value in allpeaks. Compare with tmref
                
                
                t=templatesdf[ (templatesdf["plate"]==template) & (templatesdf["well"]==w3) ]

                if (len(t)!=0):
                    #print(t.head(1)["smiles"])
                    wd["smiles"]=t.iloc[0]["smiles"]
                    wd["eos"]=t.iloc[0]["eos"]

                wd["tm"]=b[well]["tm"]

                allpeaks=[]
                if 'allpeaks' in b[well]:
                    allpeaks=b[well]['allpeaks']
                if len(allpeaks)>1:
                    wd["extra_tm_peaks"]=" ,".join(map(str,allpeaks))
                    #closest value to tm_ave
                    wd["tm"]=allpeaks[numpy.abs(numpy.array(allpeaks)-tm_ave).argmin()]

                wd["ref"]=float("{:.2f}".format(tm_ave))
                wd["ref_std"]=float("{:.2f}".format(tm_std))
                wd["dtm"]=float("{:.2f}".format(float(wd["tm"]) - tm_ave))
                wd["absdtm"]=abs(wd["dtm"])

                #If tm==0, skip.
                #print(str(wd["dtm"]))
                if (wd["tm"]==0):
                    #print("ZERO")
                    continue
                    

                #print(well,wd["tm"],tm_ave,threshold)
                if ((wd["tm"]-threshold)>tm_ave):
                    
                    wd["Effect"]="Stabilizing"
                    
                if (wd["tm"]+threshold<tm_ave):
                    
                    wd["Effect"]="Destabilizing"
                    
                #print(wd["Effect"])
                #then check what in server
                a=json.loads(serve_categories(item))
                #print(a)
                #print (well,a[well])
                
                if well in a:
                    dat=a[well]
                    #print(well,a[well])
                    #Found in persistent data
                    if ("Validated") in dat:
                        wd["Validated"]=True
                        if ("Stabilizing") in dat:
                            wd["Effect"]="Stabilizing"
                        if ("Destabilizing") in dat:
                            wd["Effect"]="Destabilizing"
                        if ("Uncertain") in dat:
                            wd["Effect"]="Uncertain"
                    if ("Warning") in dat:
                        wd["Warning"]=True
                #else:
                    #Not found, set categories from calculations.
                    #print ("NF")
                    
                report_data.append(wd)
                spk_counter=spk_counter+3
                #print(b)
    #print(report_data)

    #write headers:
    row=0
    col=0


    for header in ["File","Well","Plate template","Validated","Effect","Warning","Tm","dTm","ref Tm","ref stddev","tm plot","derivative","Multiple Tm peaks","Mol. ID.","SMILES","abs dTM"]:
                worksheet1.write(row,col,header,header_format)
                worksheet2.write(row,col,header,header_format)
                worksheet3.write(row,col,header,header_format)
                col=col+1

    row=1
    rowv=1
    rowe=1
    j={}
    worksheet1.set_column(0, 0, 40)
    worksheet2.set_column(0, 0, 40)
    worksheet3.set_column(0, 0, 40)
    worksheet4.set_column(0, 0, 40)

    worksheet1.set_column(4, 4, 12)
    worksheet2.set_column(4, 4, 12)
    worksheet3.set_column(4, 4, 12)
    
    worksheet1.set_column(2, 2, 20)
    worksheet2.set_column(2, 2, 20)
    worksheet3.set_column(2, 2, 20)

    worksheet1.set_column(12, 13, 15)
    worksheet2.set_column(12, 13, 15)
    worksheet3.set_column(12, 13, 15)

    worksheet1.set_column(14, 14, 120)
    worksheet2.set_column(14, 14, 120)
    worksheet3.set_column(14, 14, 120)

    

    for item in report_data:
        myformat=neutral_format
        if (item["Effect"]=="Stabilizing"):
            myformat=stab_format
        if (item["Effect"]=="Destabilizing"):
            myformat=destab_format
        if (item["Effect"]=="Uncertain"):
            myformat=uncert_format
            
        worksheet3.write(row,0,item["File"],myformat)
        worksheet3.write(row,1,item["Well"],myformat)
        worksheet3.write(row,3,item["Validated"],myformat)
        worksheet3.write(row,4,item["Effect"],myformat)
        worksheet3.write(row,5,item["Warning"],myformat)
        worksheet3.write(row,6,item["tm"],myformat)
        worksheet3.write(row,7,item["dtm"],myformat)
        worksheet3.write(row,8,item["ref"],myformat)
        worksheet3.write(row,9,item["ref_std"],myformat)
        worksheet3.write(row,12,item["extra_tm_peaks"],myformat)

        worksheet3.write(row,2,item["template"],myformat)
        worksheet3.write(row,13,item["eos"],myformat)
        worksheet3.write(row,14,item["smiles"],myformat)
        worksheet3.write(row,15,item["absdtm"],myformat)

        #Write sparkline data
        if (not(item["File"] in j)):
            j[item["File"]]=readtxt(data_path+item["File"])
        max_count=0
        worksheet4.write(item["spkline1"],0,item["File"])
        worksheet4.write(item["spkline1"],1,item["Well"])
        for position,point in enumerate(j[item["File"]][item["Well"]]["data"]):
            #print (item["spkline1"])
            #print(position)
            #print(point["y"])
            #print (point["data"]["y"])
            worksheet4.write(item["spkline1"],position+2,point["y"])
            max_count=position

        worksheet4.write(item["spkline2"],0,item["File"])
        worksheet4.write(item["spkline2"],1,"d"+item["Well"])
        for position,point in enumerate(j[item["File"]]["d"+item["Well"]]["data"]):
            #print (item["spkline1"])
            #print(position)
            #print(point["y"])
            #print (point["data"]["y"])
            worksheet4.write(item["spkline2"],position+2,point["y"])
            max_count=position
        #Temperatures
        worksheet4.write(item["spkline1"]-1,0,item["File"])
        worksheet4.write(item["spkline1"]-1,1,item["Well"]+" temp")
        for position,point in enumerate(j[item["File"]][item["Well"]]["data"]):
            #print (item["spkline1"])
            #print(position)
            #print(point["y"])
            #print (point["data"]["y"])
            worksheet4.write(item["spkline1"]-1,position+2,point["x"])
            max_count=position
        rang1='Data'+"!"+xlsxwriter.utility.xl_rowcol_to_cell(item["spkline1"],2)+":"+xlsxwriter.utility.xl_rowcol_to_cell(item["spkline1"],max_count+2)
        worksheet3.add_sparkline(row,10,{'range':rang1})
        rang2='Data'+"!"+xlsxwriter.utility.xl_rowcol_to_cell(item["spkline2"],2)+":"+xlsxwriter.utility.xl_rowcol_to_cell(item["spkline2"],max_count+2)
        worksheet3.add_sparkline(row,11,{'range':rang2})
       
       
        if ((item["Validated"] & (not(item["Warning"])))):
            worksheet1.add_sparkline(rowv,11,{'range':rang2})
            worksheet1.add_sparkline(rowv,10,{'range':rang1})


            worksheet1.write(rowv,0,item["File"],myformat)
            worksheet1.write(rowv,1,item["Well"],myformat)
            worksheet1.write(rowv,3,item["Validated"],myformat)
            worksheet1.write(rowv,4,item["Effect"],myformat)
            worksheet1.write(rowv,5,item["Warning"],myformat)
            worksheet1.write(rowv,6,item["tm"],myformat)
            worksheet1.write(rowv,7,item["dtm"],myformat)
            worksheet1.write(rowv,8,item["ref"],myformat)
            worksheet1.write(rowv,9,item["ref_std"],myformat)
            worksheet1.write(rowv,12,item["extra_tm_peaks"],myformat)

            worksheet1.write(rowv,2,item["template"],myformat)
            worksheet1.write(rowv,13,item["eos"],myformat)
            worksheet1.write(rowv,14,item["smiles"],myformat)
            worksheet1.write(rowv,15,item["absdtm"],myformat)
            rowv=rowv+1
        if ((item["Effect"]!="Neutral") & (not(item["Warning"]))):
            worksheet2.add_sparkline(rowe,11,{'range':rang2})
            worksheet2.add_sparkline(rowe,10,{'range':rang1})


            worksheet2.write(rowe,0,item["File"],myformat)
            worksheet2.write(rowe,1,item["Well"],myformat)
            worksheet2.write(rowe,3,item["Validated"],myformat)
            worksheet2.write(rowe,4,item["Effect"],myformat)
            worksheet2.write(rowe,5,item["Warning"],myformat)
            worksheet2.write(rowe,6,item["tm"],myformat)
            worksheet2.write(rowe,7,item["dtm"],myformat)
            worksheet2.write(rowe,8,item["ref"],myformat)
            worksheet2.write(rowe,9,item["ref_std"],myformat)
            worksheet2.write(rowe,12,item["extra_tm_peaks"],myformat)

            worksheet2.write(rowe,2,item["template"],myformat)
            worksheet2.write(rowe,13,item["eos"],myformat)
            worksheet2.write(rowe,14,item["smiles"],myformat)
            worksheet2.write(rowe,15,item["absdtm"],myformat)
            rowe=rowe+1
            
        
        row=row+1

#From worksheet 3, generate worksheets 1 and 2.



    #Report excel file has four sheets:
    #1- Validated, good, non-neutral only
    #2- good, non-neutral only
    #3- All data
    #4- Data for sparklines (corresponds to spkline1, spkline2,num_datum)
    print("Report Sent")
    report.close()
    return report_path_fil


def get_ext_files(extensions):
        all_files = []
        for ext in extensions:
            all_files.extend(Path(data_path).rglob(ext))
        return all_files

def get_cache_path(filename,store_path=cache_path,prefix="txt",rwell=""):
    #print(filename)
    #print(prefix)
    #print(store_path)
    #print(rwell)
    return store_path+prefix+"#"+filename.replace("/","##").replace("\\","##").replace(":","###")+"#"+rwell

def do_background_cache(filename):
    #print ("Preparing cache in background for "+filename)
    return readtxt(data_path+filename)

def checkline_in_file(line,file):
    f=open(file, 'r')
    fl = [line.rstrip() for line in f]
    #fl=f.readlines()
    f.close()
    if line in fl:
        return True
    return False

def readtxt(filename,rwell=""):
    #txt

    #TO DO : Move cache here instead of bering on each format.
    a={}
    cachefilename2=get_cache_path(filename,cache_path,"cache","")
    #cachefilename=cache_path+"txt#"+filename.replace("/","##")+"#"+rwell
    
    if os.path.exists(cachefilename2):
        #print("Cached file found. Serving",filename,rwell)
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return decimate(a)

    if (Path(filename).suffix==".txt"):
        file = open(filename,'r')
        line = file.readline()
        file.close()
        spline=line.split()
        #print (filename, spline)
        #print(filename)
        #print("*",spline)
        #if (len(spline)<1):
            #print("SMALL SPLINE "+filename)
        if (spline[0]=="Raw"):
            a=select_well(readlightcycler(filename),rwell)
            a["format"]=formats[0]
            #return a
        if (spline[0]=="*"):

            if (checkline_in_file("[Melt Curve Raw Data]",filename)):
            #if (spline[4]=="384-Well"):
                #[Melt Curve Raw Data]
                a=select_well(readquantstudio(filename),rwell)
                a["format"]=formats[2]
                #return a
            if (checkline_in_file("[Melt Region Temperature Data]",filename)):
            #if (spline[4]=="96well"):
                #[Melt Region Temperature Data]
                a= select_well(readquantstudio96(filename),rwell)
                a["format"]=formats[3]
                #return a
    #xlsx biorad
    if (Path(filename).suffix==".xlsx"):
        #print("XLS "+filename)
        a= select_well(readbiorad(filename),rwell)
        a["format"]=formats[1]
        #return a
    #Generic
        #xlsx
    if (Path(filename).suffix==".gdsf"):
        #print("XLS "+filename)
        a= select_well(readgeneric(filename),rwell)
        a["format"]=formats[4]
        #return a
    
    f_cached = open(cachefilename2,'w')
    f_cached.write(json.dumps(a))
    f_cached.close()
    return decimate(a)


def remove_data(jsondata):
    
    #Remove data from jsondata

    for w in jsondata:
        if (w!="format"):
            jsondata[w]['data']=[]
    return jsondata
    

def select_well(jsondata,rwell=""):
    
    if rwell!="":
        jsondata={k:v for k,v in jsondata.items() if k in [rwell,'d'+rwell,'format']}
    return jsondata
    

def readtxttm(filename,rwell=""):
    #txt
    #formats=["LightCycler","BioRad","QuantStudio 384","QuantStudio 96"]
    return remove_data(readtxt(filename,rwell))


def zipdir(path, ziph):
    # ziph is zipfile handle
    for root, dirs, files in os.walk(path):
        for file in files:
            #print(os.path.relpath(os.path.join(root, file)))
            ziph.write(os.path.join(root, file),os.path.relpath(os.path.join(root, file),os.path.join(path, '.')))

def readquantstudio(filename):
    #Reads txt file and returns the graphs for each well ('A1'), and the 'gradient' ('dA1')
    #To convert to json to be used by the web client.
    #rwell=""
    #Make faster usign same trick as in biorad: First time not cached write entire dataset, then filter.



        #As this is very slow, also try to load the nowell version, then filter.
    cachefilename2=get_cache_path(filename,cache_path,"txt","")
    #cachefilename=cache_path+"txt#"+filename.replace("/","##")+"#"+rwell
    
    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a


    filter=0
    # Open the file 
    file = open(filename,'r')
    lines = file.readlines()
    file.close()
    returndata={}
    dats_for_ders={}
    #well = ''
    #new_file = ''
    firstiter=1
    #first generate data scaffold.
    for a in range(1,25):
        for b in ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P"]:
            well=b+str(a)
            returndata[well]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
            returndata["d"+well]={'data':[],'max':0,'min':0}
            dats_for_ders[well]={'d':[],'t':[]}

    for line in lines[2:]:
        spline= line.split('\t')
        if (spline[0]=="Well"):
            firstiter=0
            continue
        if (firstiter==1):
            continue
        #Here we start parsing data.
        w=spline[1]
        t=float(spline[3])
        v=float(spline[4].replace(",",""))
        #print(w,t,v)
        returndata[w]['data'].append({'x':t,'y':v})
        dats_for_ders[w]['d'].append(v)
        dats_for_ders[w]['t'].append(t)
    
    #Now we do the derivatives.
    returndata=derivate_and_metadata(returndata)

    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))

    return returndata

def readgeneric(filename,rwell=""):
    #Generic file. Three columns, separated by tabs or spaces, no header
    #Column order: Well, temperature(C), fluorescence
    cachefilename2=get_cache_path(filename,cache_path,"gdsf","")
    
    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a

    file = open(filename,'r')
    lines = file.readlines()
    file.close()
    returndata={}
    dats_for_ders={}
    #first generate data scaffold.
    for a in range(1,25):
        for b in ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P"]:
            well=b+str(a)
            returndata[well]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
            returndata["d"+well]={'data':[],'max':0,'min':0}
            dats_for_ders[well]={'d':[],'t':[]}


    for line in lines:
        spline= line.split()
        well=spline[0]
        temp=spline[1]
        fluor=spline[2]
        returndata[well]['data'].append({'x':float(temp),'y':float(fluor)})

    #rint(returndata)
    returndata=derivate_and_metadata(returndata)

    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))
    

    #print("*",returndata)
    return returndata


def readquantstudio96(filename,rwell=""):
    #Reads txt file and returns the graphs for each well ('A1'), and the 'gradient' ('dA1')
    #To convert to json to be used by the web client.
    #rwell=""
    #Make faster usign same trick as in biorad: First time not cached write entire dataset, then filter.


    
    cachefilename2=get_cache_path(filename,cache_path,"txt","")
    
    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a
    filter=0
    # Open the file 
    file = open(filename,'r')
    lines = file.readlines()
    file.close()
    returndata={}
    dats_for_ders={}
    #well = ''
    #new_file = ''
    phase=0
    #first generate data scaffold.
    for a in range(1,25):
        for b in ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P"]:
            well=b+str(a)
            returndata[well]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
            returndata["d"+well]={'data':[],'max':0,'min':0}
            dats_for_ders[well]={'d':[],'t':[]}


#96 well format is split in to different blocks.
#first, after we find "[Melt Region Temperature Data]", we skip one lines.
#Then we split by tab: [1]->Well [4]->[End]: Temperatures
#We skip blank lines
#And then we find "[Melt Region Normalized Data]". We skip one line
#Then we split by tab: [1]->Well [4]->[End]: Fluorescence
    for line in lines[2:]:
        spline= line.split('\t')
        #print("phase:",phase)
        #print(spline[0])
        if (spline[0]=="[Melt Region Normalized Data]\n"):
            phase=phase+1
            continue
        if (spline[0]=="[Melt Region Temperature Data]\n"):
            phase=phase+1
            continue

        if (phase==1):
            phase=phase+1
            continue
        if (phase==4):
            phase=phase+1
            continue
        if (phase==2):
            if (len(spline)<2):
                #Newline:We found the end of temperatures block
                phase=phase+1
                continue
            #Here we start parsing data:Temperature.
            w=spline[1]
            tblock=spline[4:]
            for i in tblock:
                t=float(i)
                returndata[w]['data'].append({'x':t,'y':0})
                dats_for_ders[w]['t'].append(t)

        if (phase==5):
            if (len(spline)<2):
                #Newline:We found the end of fluorescence block
                phase=phase+1
                continue
            #Here we start parsing data:fluorescence.
            w=spline[1]
            fblock=spline[4:]
            counter=0
            for i in fblock:
                
                f=float(i)
                returndata[w]['data'][counter]['y']=f
                dats_for_ders[w]['d'].append(f)
                counter=counter+1
        #Everything is read after phase 5

        
    #Now we do the derivatives.
    returndata=derivate_and_metadata(returndata)

    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))
    

    #print("*",returndata)
    return returndata

def derivate_and_metadata(jsondata):
    #calculates derivatives, maxs, mins and tms.
    dats_for_ders={}
    for well in jsondata:
        if well[0]!="d":
            #print(well)
            dats_for_ders[well]={'d':[],'t':[]}
            for point in jsondata[well]['data']:
                #print(point)
                #print(point,point['y'])
                dats_for_ders[well]['d'].append(point['y'])
                dats_for_ders[well]['t'].append(point['x'])
            fluors=dats_for_ders[well]['d']
            temps=dats_for_ders[well]['t']
            
            if (len(fluors)>0):
                dfluors=numpy.gradient(fluors)
                dfluors=savgol_filter(dfluors,31,2)
                
                
                dfluors_poly = resample_poly(dfluors, 100, 20)
                min_temp_v=temps[0]
                max_temp_v=temps[-1]
                new_temps=np.linspace(min_temp_v,max_temp_v,len(dfluors_poly))


                #maxes_der=argrelextrema(dfluors,numpy.greater,order=10)
                
                fmax=numpy.max(fluors)
                fmin=numpy.min(fluors)

                dfmax=numpy.max(dfluors)
                dfmin=numpy.min(dfluors)


                maxes_der,maxes_properties=find_peaks(dfluors_poly, prominence=0.3*(dfmax-dfmin))
                for di in range(0,len(dfluors)):
                        ddataitem={'x':temps[di],'y':dfluors[di]}
                        jsondata["d"+well]['data'].append(ddataitem)
                
                #tm=temps[numpy.argmax(dfluors)]
                tm=round(new_temps[numpy.argmax(dfluors_poly)],2)
                jsondata[well]['tm']=tm
                jsondata[well]['max']=fmax
                jsondata[well]['min']=fmin 
                jsondata['d'+well]['max']=dfmax
                jsondata['d'+well]['min']=dfmin
                tm_array=[]
                for mi in maxes_der:
                    #tm_array.append(temps[mi])
                    tm_array.append(round(new_temps[mi],2))
                #if len(tm_array)!=1:
                    #print(well,tm,"-",tm_array)
                if not(tm in tm_array):
                    tm_array.append(tm) 
                jsondata[well]['allpeaks']=tm_array
                warning=checkquality(jsondata[well]['data'],jsondata['d'+well]['data'])
                jsondata[well]['warnings']=warning


                
    return jsondata

def decimate(jsondata):


    #print(jsondata)
    #Decimate data: Return a max of 100 points of the data in each well and dwell.
    
    if ( not decimation):
        return jsondata

    #print ("DECIM ON")
    for well in jsondata:
        #print (well)
        
        #print(jsondata[well])
        if 'data' in jsondata[well]:
            #print (len(jsondata[well]['data']))
            numpoints=len(jsondata[well]['data'])
            if (numpoints<=decim_maxpoints):
                continue
            point_every=math.floor(float(numpoints)/float(decim_maxpoints))
            w=[]
            count=0
            for point in jsondata[well]['data']:
                count=count+1
                if (count % point_every==0):
                    w.append(point)
            #print(jsondata[well]['data'])
            #print(w)
            jsondata[well]['data']=w

    return jsondata

def readbiorad(filename,rwell=""):



    cachefilename2=get_cache_path(filename,cache_path,"xlsx","")
    #cachefilename=cache_path+"txt#"+filename.replace("/","##")+"#"+rwell

    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a

    #This gives errors because data is not really excel compliant.
    #Zip is malformed,a sseparators are \ instead of /

    returndata={}    
    #First, fix separators.
    z=zipfile.ZipFile(filename)
    extract_path=tempfile.mkdtemp()
    #print(extract_path)
    for i, f in enumerate(z.filelist):
        f.filename=f.filename.replace('\\','/')
        if (f.filename=="[content_types].xml"):
            f.filename='[Content_Types].xml'
        if (f.filename=="xl/sharedstrings.xml"):
            f.filename='xl/sharedStrings.xml'

        z.extract(f,path=extract_path)
    #And save file
    z.close()
    xlsxpath=tempfile.mkdtemp()
    fixed_filename=xlsxpath+'/fixed.xlsx'
    zipf = zipfile.ZipFile(fixed_filename, 'w', zipfile.ZIP_DEFLATED)
    zipdir (extract_path+"/",zipf)
    zipf.close()
    #print(fixed_filename)

    #Now it should be fixed.
    df=pandas.read_excel(fixed_filename)
    #print(df)
    columns=df.columns[2:]
    #print(columns)

    #first generate data scaffold.
    for a in range(1,25):
        for b in ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P"]:
            well=b+str(a)
            returndata[well]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
            returndata["d"+well]={'data':[],'max':0,'min':0}
            #dats_for_ders[well]={'d':[],'t':[]}



    #For each column
    for w in columns:
        dataitem={}
        ddataitem={}
        temps=[]
        fluors=[]
        #returndata[w]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
        #returndata["d"+w]={'data':[],'max':0,'min':0}
        dat=df[['Temperature',w]]
        for index,row in dat.iterrows():
            temps.append(float(row["Temperature"]))
            fluors.append(float(row[w]))
        #print(w,temps,fluors)
        dfluors=[]
        for i in range(0,len(temps)):
            dataitem={'x':temps[i],'y':fluors[i]}
            returndata[w]['data'].append(dataitem)
            #dfluor=0

    #We have all data here. Now do derivatives and metadata.
    returndata=derivate_and_metadata(returndata)


    #Save all data
    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))
    #f_cached.close()
    

        
    return returndata
    


def readbiorad_old(filename,rwell=""):



    cachefilename2=get_cache_path(filename,cache_path,"xlsx","")
    #cachefilename=cache_path+"txt#"+filename.replace("/","##")+"#"+rwell

    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a

    #This gives errors because data is not really excel compliant.
    #Zip is malformed,a sseparators are \ instead of /

    returndata={}    
    #First, fix separators.
    z=zipfile.ZipFile(filename)
    extract_path=tempfile.mkdtemp()
    #print(extract_path)
    for i, f in enumerate(z.filelist):
        f.filename=f.filename.replace('\\','/')
        if (f.filename=="[content_types].xml"):
            f.filename='[Content_Types].xml'
        if (f.filename=="xl/sharedstrings.xml"):
            f.filename='xl/sharedStrings.xml'

        z.extract(f,path=extract_path)
    #And save file
    z.close()
    xlsxpath=tempfile.mkdtemp()
    fixed_filename=xlsxpath+'/fixed.xlsx'
    zipf = zipfile.ZipFile(fixed_filename, 'w', zipfile.ZIP_DEFLATED)
    zipdir (extract_path+"/",zipf)
    zipf.close()
    #print(fixed_filename)

    #Now it should be fixed.
    df=pandas.read_excel(fixed_filename)
    #print(df)
    columns=df.columns[2:]
    #print(columns)



    #For each column
    for w in columns:
        dataitem={}
        ddataitem={}
        temps=[]
        fluors=[]
        returndata[w]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
        returndata["d"+w]={'data':[],'max':0,'min':0}
        dat=df[['Temperature',w]]
        for index,row in dat.iterrows():
            temps.append(float(row["Temperature"]))
            fluors.append(float(row[w]))
        #print(w,temps,fluors)
        dfluors=[]
        for i in range(0,len(temps)):
            dataitem={'x':temps[i],'y':fluors[i]}
            returndata[w]['data'].append(dataitem)
            #dfluor=0
            if (i>0):
                dfluors.append(fluors[i]-fluors[i-1])
                
        #The original software smoothed the derivative  using a savitzky-golay filter (sg_filter.py).
        #Here, smooth the derivative using savitzky-golag (sci-py)
        dfluors=savgol_filter(dfluors,31,2)
        for i in range(0,len(dfluors)):
            ddataitem={'x':temps[i+1],'y':dfluors[i]}
            returndata["d"+w]['data'].append(ddataitem)
        tm=temps[numpy.argmax(dfluors)+1]
        returndata[w]['tm']=tm
        fmax=numpy.max(fluors)
        fmin=numpy.min(fluors)
        returndata[w]['max']=fmax
        returndata[w]['min']=fmin        
        dfmax=numpy.max(dfluors)
        dfmin=numpy.min(dfluors)
        returndata["d"+w]['max']=dfmax 
        returndata["d"+w]['min']=dfmin

    #Save all data
    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))
    #f_cached.close()
    

        
    return returndata

def readlightcycler(filename,rwell=""):

    #Reads txt file and returns the graphs for each well ('A1'), and the 'gradient' ('dA1')
    #To convert to json to be used by the web client.
    #rwell=""
    #Make faster usign same trick as in biorad: First time not cached write entire dataset, then filter.


        #As this is very slow, also try to load the nowell version, then filter.
    cachefilename2=get_cache_path(filename,cache_path,"txt","")
    #cachefilename=cache_path+"txt#"+filename.replace("/","##")+"#"+rwell

    if os.path.exists(cachefilename2):
        #print("Cached file no-well found. Serving")
        f_served = open(cachefilename2,'rb')
        a=json.loads(f_served.read())
        return a





    #Now fill wells data

    filter=0
    # Open the file 
    file = open(filename,'r')
    lines = file.readlines()
    file.close()
    returndata={}
    dats_for_ders={}

    #first generate data scaffold.
    for a in range(1,25):
        for b in ["A","B","C","D","E","F","G","H","I","J","K","L","M","N","O","P"]:
            well=b+str(a)
            returndata[well]={'data':[],'tm':0,'warnings':0,'max':0,'min':0}
            returndata["d"+well]={'data':[],'max':0,'min':0}
            dats_for_ders[well]={'d':[],'t':[]}

    for line in lines[2:]:
        spline = line.split('\t')
        #Fix the well name
        #columns:number
        #row:letter
        
        w = spline[0]
        row=well[0]
        column=well[1:]
        t=float(spline[6])
        v=float(spline[7])
        returndata[w]['data'].append({'x':t,'y':v})
        dats_for_ders[w]['d'].append(v)
        dats_for_ders[w]['t'].append(t)

    #All data read.
    
    #Now we do the derivatives.
    returndata=derivate_and_metadata(returndata)
    #f_cached = open(cachefilename2,'w')
    #f_cached.write(json.dumps(returndata))


    #print(returndata)
    #f_cached = open(cachefilename,'w')
    #f_cached.write(json.dumps(returndata))
    return returndata



def checkquality(qdat,dqdat):

    #Checks data quality, returns warnings from 0:Ok, to max:Bad
    warning=0
    #I disable dthe autocheck because it does not work very well.
    return(warning)

    
    #Check 1.
    #First quarter > last quarter.
    
    
    a=len(qdat)
    v1=0
    v4=0
    for i in range(0,int(float(a)/4)):
        v1=v1+qdat[i]['y']
    for i in range( int (( float(a)/4)*3) ,a):
        v4=v4+qdat[i]['y']
    #print(v1,v4)
    if (v1>v4):
        warning=warning+1
    return warning

def store_categories(filename,rwell,categories):
    
        categoriesfilename=get_cache_path(filename,persistent_path,"cat","")
        #categoriesfilename=persistent_path+"cat#"+filename.replace("/","##")
        #rwell
        
        
        cats={}
        #If file exists, read it.
        if os.path.exists(categoriesfilename):
            #print("Categories file found. Reading")
            f = open(categoriesfilename,'rb')
            cats= json.loads(f.read())
            f.close()
        f_cat = open(categoriesfilename,'w')
        cats[rwell]=categories
        


        f_cat.write(json.dumps(cats))
    
def serve_categories(filename):
    
        categoriesfilename=get_cache_path(filename,persistent_path,"cat","")
        #categoriesfilename=persistent_path+"cat#"+filename.replace("/","##")
        if os.path.exists(categoriesfilename):
            #print("Categories file found. Serving")
            f_served = open(categoriesfilename,'r')
            #return json.loads(f_served.read())
            return f_served.read()
        else:
            return "{}"

def store_template(filename,template):
    
        #template for each filename are stored at persistent_path/templates.txt
        templatesfilename=persistent_path+ os.path.sep+"templates.txt"
        #print(templatesfilename,filename,template)
        #categoriesfilename=persistent_path+"cat#"+filename.replace("/","##")
        a={}
        if os.path.exists(templatesfilename):
            #print("Categories file found. Serving")
            f_served = open(templatesfilename,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_served.read())

        
        a[filename]=template
        f_temp = open(templatesfilename,'w')
        f_temp.write(json.dumps(a))

def serve_template(filename):
    
        #template for each filename are stored at persistent_path/templates.txt
        templatesfilename=persistent_path+ os.path.sep+"templates.txt"
        #print(templatesfilename,filename)
        #categoriesfilename=persistent_path+"cat#"+filename.replace("/","##")
        if os.path.exists(templatesfilename):
            #print("Categories file found. Serving")
            f_served = open(templatesfilename,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_served.read())
            if filename in a:
                #print(a[filename])
                return a[filename]
            else:
                #print("None")
                return "None"

        else:
            return "None"


def store_kd_template(data):
    #print(data)
    d=json.loads(data)
    kd_templatefile=persistent_path+"kdtemplates.templ"
    a=[]
    if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            f_templ.close()
    
    
    a.append(d)
    f_templ = open(kd_templatefile,'w')
    f_templ.write(json.dumps(a))

def serve_kd_templatelist():

        #print("Asked for templates")
        kd_templatefile=persistent_path+"kdtemplates.templ"
        templates=[]
        if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            for templs in a:
                templates.append(templs["filename"])

        #get unique values 
        templates= list(set(templates))
        return str.encode(json.dumps(templates))

def serve_templatelist():

        result=[]
        templatesdf=[]
        #for path in Path(data_path).rglob(('*.txt','*.xlsx')):
        extensions=['*.txt']
        all_files = []
        for ext in extensions:
            all_files.extend(Path(plateinfo_path).rglob(ext))
        for filename in all_files:
                #print(filename)
                df = pandas.read_csv(filename, index_col=None, sep="\t",header=0)
                templatesdf.append(df)
        templatesdf=pandas.concat(templatesdf, axis=0, ignore_index=True)
        templatesdf.columns=["platewell","eos","smiles","plate","well"]
        
        result=templatesdf["plate"].unique()
        return str.encode(json.dumps(result.tolist()))

def get_filelist_for_cache():


        result=[]
        #for path in Path(data_path).rglob(('*.txt','*.xlsx')):
        for path in get_ext_files(extension_list):
                #print(path)
                filenam=os.path.relpath(path,data_path)
                filenams=filenam.split(os.path.sep)
                if (filenams[-1][:1]=="."):
                    continue
                result.append(filenam)
        return result


class SimpleHTTPRequestHandler(BaseHTTPRequestHandler):

    def get_kd_file(self,filename):
        response={}
        #template for each filename are stored at persistent_path/templates.txt
        kd_templatefile=persistent_path+"kdtemplates.templ"
        #print(templatesfilename,filename)
        #categoriesfilename=persistent_path+"cat#"+filename.replace("/","##")
        if os.path.exists(kd_templatefile):
            #print("Categories file found. Serving")
            f_templ = open(kd_templatefile,'r')
            #return json.loads(f_served.read())
            a=json.loads(f_templ.read())
            for item in a:
                if item["filename"]==filename:
                    response=item
                    
        
        return str.encode(json.dumps(response))
 
    def get_datafile(self,filename):
        
        fil=data_path+filename
        #print("Asked for: "+fil)
        #Read filename and return it as:
        '''
        [cell:{column:'A',row:'1'}, acq:[ data:[{x:25,y:7000},{x:25,y:7000},...], deriv:[] ]
        
        '''
        #It needs to be possible to extract acq and deriv and send it to Chart.js directly.
        #https://www.chartjs.org/docs/latest/general/data-structures.html
        a=readtxt(fil)
        #return str.encode(json.dumps([{'x':1,'y':1},{'x':2,'y':2},{'x':3,'y':6},{'x':4,'y':2},{'x':5,'y':1}]))
        return str.encode(json.dumps(a))

    def get_tms_for_plate(self,filename):
        
        fil=data_path+filename
        #print("Asked for Tms: "+fil)
        #Read filename and return it as:
        '''
        [cell:{column:'A',row:'1'}, acq:[ data:[{x:25,y:7000},{x:25,y:7000},...], deriv:[] ]
        
        '''
        #It needs to be possible to extract acq and deriv and send it to Chart.js directly.
        #https://www.chartjs.org/docs/latest/general/data-structures.html
        a=readtxttm(fil)
        #return str.encode(json.dumps([{'x':1,'y':1},{'x':2,'y':2},{'x':3,'y':6},{'x':4,'y':2},{'x':5,'y':1}]))
        return str.encode(json.dumps(a))

    def get_well(self,filename,rwell=""):
        
        fil=data_path+filename
        #print("Asked for: "+fil)
        #Read filename and return it as:
        '''
        [cell:{column:'A',row:'1'}, acq:[ data:[{x:25,y:7000},{x:25,y:7000},...], deriv:[] ]
        
        '''
        #It needs to be possible to extract acq and deriv and send it to Chart.js directly.
        #https://www.chartjs.org/docs/latest/general/data-structures.html
        a=readtxt(fil,rwell)
        #return str.encode(json.dumps([{'x':1,'y':1},{'x':2,'y':2},{'x':3,'y':6},{'x':4,'y':2},{'x':5,'y':1}]))
        return str.encode(json.dumps(a))




    def log_message(self, format, *args):
        #Print to log file
        return



    def get_filelist(self):
        result=[]
        #for path in Path(data_path).rglob(('*.txt','*.xlsx')):
        for path in get_ext_files(extension_list):
                #print(path)
                filenam=os.path.relpath(path,data_path)
                result.append(filenam)
        return str.encode(json.dumps(result))

    def get_filelist_nested(self):
        result={}
        for path in sorted(get_ext_files(extension_list)):
        #for path in sorted(Path(data_path).rglob(('*.txt','*.xlsx'))):
                #print(path)
                filenam=os.path.relpath(path,data_path)
                filenams=filenam.split(os.path.sep)
                p=filenams[0]
                #print(p,filenam,filenams,filenams[-1][:1])
                #Ignore unix hiddenfiles
                if (filenams[-1][:1]=="."):
                    continue
                fil=os.path.sep.join(filenams[1:])
                r=result.get(p,[])
                r.append(fil)
                result[p]=r
                #result.append(filenam)
        #print(result)
        return str.encode(json.dumps(result))


    def do_POST(self):
        splitpath=self.path.split("/")

        if (self.path=="/save_kd_template/"):

            content_length = int(self.headers['Content-Length']) # <--- Gets the size of data
            post_data = self.rfile.read(content_length) # <--- Gets the data itself
            self.send_response(200)
            self.end_headers()
            store_kd_template(string_escape(post_data.decode('utf-8')))
            self.wfile.write(str.encode("OK"))
        
        
        elif (splitpath[1]=="get_report"):

            content_length = int(self.headers['Content-Length']) # <--- Gets the size of data
            post_data = self.rfile.read(content_length) # <--- Gets the data itself
            #print("POST request,\nPath: %s\nHeaders:\n%s\n\nBody:\n%s\n",
                    #str(self.path), str(self.headers), post_data.decode('utf-8'))

            self.send_response(200)
            self.end_headers()
            
            #Excel file at tmp

            content_path=create_report(post_data.decode('utf-8'))
            with open(content_path, 'rb') as content:
                shutil.copyfileobj(content, self.wfile)
            #self.wfile.write("POST request for {}".format(self.path).encode('utf-8'))
        elif (splitpath[1]=="get_kd_report"):

            content_length = int(self.headers['Content-Length']) # <--- Gets the size of data
            post_data = self.rfile.read(content_length) # <--- Gets the data itself
            #print("POST request,\nPath: %s\nHeaders:\n%s\n\nBody:\n%s\n",
                    #str(self.path), str(self.headers), post_data.decode('utf-8'))

            self.send_response(200)
            self.end_headers()
            
            #Excel file at tmp

            content_path=create_kd_report(post_data.decode('utf-8'))
            with open(content_path, 'rb') as content:
                shutil.copyfileobj(content, self.wfile)
            #self.wfile.write("POST request for {}".format(self.path).encode('utf-8'))
        elif (splitpath[1]=="get_kd_data"):

            content_length = int(self.headers['Content-Length']) # <--- Gets the size of data
            post_data = self.rfile.read(content_length) # <--- Gets the data itself
            #print("POST request,\nPath: %s\nHeaders:\n%s\n\nBody:\n%s\n",
                    #str(self.path), str(self.headers), post_data.decode('utf-8'))

            self.send_response(200)
            self.end_headers()
            
            #return json data
            f_content=str.encode(json.dumps(get_kd_data(post_data.decode('utf-8'))))
            self.wfile.write(f_content)
        elif (splitpath[1]=="get_kd_data_molecule"):

            content_length = int(self.headers['Content-Length']) # <--- Gets the size of data
            post_data = self.rfile.read(content_length) # <--- Gets the data itself
            #print("POST request,\nPath: %s\nHeaders:\n%s\n\nBody:\n%s\n",
                    #str(self.path), str(self.headers), post_data.decode('utf-8'))

            self.send_response(200)
            self.end_headers()
            
            #return json data
            f_content=str.encode(json.dumps(get_kd_data_m(post_data.decode('utf-8'))))
            self.wfile.write(f_content)


    def do_GET(self):
        #print(self.path)
        splitpath=self.path.split("/")
        #print(splitpath[1])

        #requesting a list of files
        #This is a list of files from the data_path
        if (self.path=="/filelist/"):
            self.send_response(200)
            self.end_headers()
            f_content=self.get_filelist()

        elif (self.path=="/filelistn/"):
            self.send_response(200)
            self.end_headers()
            f_content=self.get_filelist_nested()
        
        elif (self.path=="/templatelist/"):
            self.send_response(200)
            self.end_headers()
            f_content=serve_templatelist()

        elif (self.path=="/kdtemplatelist/"):
            self.send_response(200)
            self.end_headers()
            f_content=serve_kd_templatelist()

        elif (splitpath[1]=="well"):
            self.send_response(200)
            self.end_headers()
            #print(self.path)
            req_filename=unquote("/".join(splitpath[2:]))
            #print(req_filename)
            splf=req_filename.split("&")
            req_filename=splf[0]
            req_well=splf[1]
            #print("*"+req_filename+"*"+req_well)
            #rew_well=splitpath
            f_content=self.get_well(req_filename,req_well)

        #requesting a data file
        elif (splitpath[1]=="file"):
            self.send_response(200)
            self.end_headers()
            req_filename=unquote("/".join(splitpath[2:]))
            #print("*"+req_filename)
            f_content=self.get_datafile(req_filename)

        #requesting a kd_template file
        elif (splitpath[1]=="kd_template"):
            self.send_response(200)
            self.end_headers()
            req_filename=unquote("/".join(splitpath[2:]))
            #print("*"+req_filename)
            f_content=self.get_kd_file(req_filename)

        #requesting a Tms
        elif (splitpath[1]=="tm"):
            self.send_response(200)
            self.end_headers()
            req_filename=unquote("/".join(splitpath[2:]))
            #print("*"+req_filename)
            f_content=self.get_tms_for_plate(req_filename)
        #Receive categories from client
        elif (splitpath[1]=="send_categories"):
            self.send_response(200)
            self.end_headers()
            args="/".join(splitpath[2:])
            splf=args.split("&")
            filename=unquote(splf[0])
            well=splf[1]
            categories=splf[2].split("$")

            #print(filename,well,categories)
            store_categories(filename,well,categories)
            f_content=str.encode("OK")

        elif (splitpath[1]=="get_categories"):
            self.send_response(200)
            self.end_headers()
            #print(self.path)
            req_filename=unquote("/".join(splitpath[2:]))
            #print(req_filename)
            splf=req_filename.split("&")
            req_filename=splf[0]
            #req_well=splf[1]
            #print("*"+req_filename+"*"+req_well)
            #rew_well=splitpath
            f_content=str.encode(serve_categories(req_filename))

        elif (splitpath[1]=="get_template"):
            self.send_response(200)
            self.end_headers()
            #print(self.path)
            req_filename=unquote("/".join(splitpath[2:]))
            #print(req_filename)
            splf=req_filename.split("&")
            req_filename=splf[0]
            #req_well=splf[1]
            #print("*"+req_filename+"*"+req_well)
            #rew_well=splitpath
            f_content=str.encode(serve_template(req_filename))

        elif (splitpath[1]=="send_template"):
            self.send_response(200)
            self.end_headers()
            args="/".join(splitpath[2:])
            splf=args.split("&")
            filename=unquote(splf[0])
            template=splf[1]
            
            store_template(filename,template)
            f_content=str.encode("OK")
        
        
        
        #general case
       
       
       
        else:
            self.send_response(200)
            self.end_headers()
            # place absolute path here
            try:
                file_p=serv_path+self.path
                #print(file_p)
                f_served = open(file_p,'rb')
                f_content = f_served.read()
            except IOError:
                #print("File not accessible")
                f_content=str.encode("")
            #finally:
                #f_served.close()
        self.wfile.write(f_content)


if __name__ == "__main__":

        freeze_support() # Enable multi-threading for windows pyinstaller version
        #Read settings file for paths.
        extension_list=('*.txt', '*.xlsx','*.gdsf')
        settingsfile="settings.ini"
        if (os.path.exists(settingsfile)):
            #open and read it. If not found, use defaults.
            config=configparser.ConfigParser()
            config.read(settingsfile)
            if ('Default' in config):
                    if ('data_path' in config['Default']):
                        data_path=config['Default']['data_path']
                        print("data_path set to "+data_path)
                    if ('cache_path' in config['Default']):
                        cache_path=config['Default']['cache_path']
                        print("cache_path set to "+cache_path)
                    if ('persistent_path' in config['Default']):
                        persistent_path=config['Default']['persistent_path']
                        print("persistent_path set to "+persistent_path)
                    if ('plateinfo_path' in config['Default']):
                        plateinfo_path=config['Default']['plateinfo_path']
                        print("plateinfo_path set to "+plateinfo_path)

        #Create paths if not there:
        paths=[data_path,cache_path,persistent_path,plateinfo_path]
        for mypath in paths:

                if not(os.path.exists(mypath)):
                    print("Creating folder " + mypath)
                    Path(mypath).mkdir(parents=True, exist_ok=True)
                    
                
            

        
        docache=True
        
        parser = argparse.ArgumentParser(description='Server for DSF HTS.')
        parser.add_argument('--quick', action='store_true', help='Do not build cache at start')
        args = parser.parse_args()
        if (args.quick):
            docache=False
        #docache=False
        #check old version
        is_new_version=False
        old_version="Unknown"
        if (os.path.exists(persistent_path+"version.txt")):
            
                f = open(persistent_path+"version.txt", "r")
                a=f.readline()
                #print(a,version)
                old_version=a.rstrip()
                if not (a==version):
                    is_new_version=True

            
        else:
            
                is_new_version=True
            
        if (is_new_version):
            print("New version: "+version+" (Old version:"+old_version+"). Removing cache.")
            f = open(persistent_path+"version.txt", "w")
            f.write(version)
            f.close()
            remove_cache()

        httpd = HTTPServer(('localhost', 5555), SimpleHTTPRequestHandler)

        if docache:
            print("Please wait. Scanning files.")
            fl=get_filelist_for_cache()
            #print (fl)
            with Pool(1) as p:
                p.map(do_background_cache,fl)
        
        print("PROGRAM STARTED. Open:")
        myurl="http://localhost:5555/index.html"
        print(myurl)
        webbrowser.open(myurl)
        httpd.serve_forever()



#pyinstaller --add-data=webapp;webapp --add-data=settings.ini;. server.py
